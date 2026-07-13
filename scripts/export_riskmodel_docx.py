"""
Exportiert docs/risk_model.md als formatiertes Word-Dokument.
Output: docs/risk_model.docx
"""

from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# ── Seitenränder ─────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

# ── Hilfsfunktionen ───────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)

def add_heading(text, level=1):
    colors = {1: "1A237E", 2: "283593", 3: "37474F"}
    sizes  = {1: 16, 2: 13, 3: 11}
    p = doc.add_paragraph()
    if level == 1:
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after  = Pt(6)
    else:
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(sizes[level])
    run.font.color.rgb = RGBColor.from_string(colors[level])
    return p

def add_para(text="", bold_parts=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    if bold_parts:
        for part, is_bold in bold_parts:
            run = p.add_run(part)
            run.bold = is_bold
            run.font.size = Pt(10.5)
    else:
        run = p.add_run(text)
        run.font.size = Pt(10.5)
    return p

def add_code(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.font.name = "Courier New"
    run.font.size = Pt(9.5)
    run.font.color.rgb = RGBColor(0x1A, 0x23, 0x7E)
    return p

def add_table(headers, rows, col_widths, header_color="37474F", row_colors=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Header
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = h
        set_cell_bg(cell, header_color)
        run = cell.paragraphs[0].runs[0]
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(10)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Datenzeilen
    for r_idx, row_data in enumerate(rows):
        row = table.rows[r_idx + 1]
        bg = (row_colors[r_idx] if row_colors and r_idx < len(row_colors) else None)
        for c_idx, val in enumerate(row_data):
            cell = row.cells[c_idx]
            cell.text = str(val)
            cell.paragraphs[0].runs[0].font.size = Pt(10)
            if bg:
                set_cell_bg(cell, bg)

    # Spaltenbreiten
    for i, w in enumerate(col_widths):
        for row in table.rows:
            row.cells[i].width = Cm(w)

    doc.add_paragraph()


# ══════════════════════════════════════════════════════════════════════════════
#  INHALT
# ══════════════════════════════════════════════════════════════════════════════

# Titel
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Risikomodell — SC Risk System")
run.bold = True
run.font.size = Pt(20)
run.font.color.rgb = RGBColor(0x1A, 0x23, 0x7E)

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = p2.add_run("Stand: 2026-06-01")
r2.font.size = Pt(10)
r2.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

doc.add_paragraph()

# ── Was berechnen wir? ────────────────────────────────────────────────────────
add_heading("Was berechnen wir?", 1)
add_para("Für jede betroffene Fabrik (Facility) berechnen wir einen RiskScore (0–100). "
         "Ein hoher Score bedeutet: Diese Fabrik ist sehr gefährdet, und es gibt kaum Alternativen.")
add_para("Die drei Fabriken mit dem höchsten Score erscheinen im Risikobericht.")

# ── Hauptformel ───────────────────────────────────────────────────────────────
add_heading("Die Hauptformel", 1)
add_code("RiskScore = Severity  ×  TierWeight  ×  Vulnerability  ×  (1 − ResilienceDiscount)")
add_para("Danach wird normalisiert:")
add_code("RiskScore (0–100)  =  RiskScore_roh  /  5.0  ×  100")
add_para("Das Maximum ist 5.0 (Severity=5, TierWeight=1.0, Vulnerability=1.0, kein Discount).")

# ── Schritt 1: Severity ───────────────────────────────────────────────────────
add_heading("Schritt 1 — Severity", 1)
add_para("Wie schlimm ist das Ereignis? Das LLM liest den Nachrichtentext "
         "(Risk Assessment Agent) und gibt eine Zahl von 1 bis 5 zurück.")

add_table(
    headers=["Wert", "Bedeutung", "Beispiel"],
    rows=[
        ["1", "Sehr gering", "Kleiner lokaler Streik, nur 1 Tag"],
        ["2", "Gering",      "Kurze Lieferverzögerung"],
        ["3", "Mittel",      "Monatelanger Streik"],
        ["4", "Hoch",        "Exportverbot für kritisches Material"],
        ["5", "Kritisch",    "Kompletter Lieferstopp, keine Alternative"],
    ],
    col_widths=[2, 4, 9],
    row_colors=["DCEDC8","F9FBE7","FFF9C4","FFE0B2","FFCDD2"],
)

# ── Schritt 2: TierWeight ─────────────────────────────────────────────────────
add_heading("Schritt 2 — TierWeight", 1)
add_para(
    "TierWeight hängt davon ab, wie weit eine Fabrik vom Ereignisursprung (origin_tier) entfernt ist. "
    "Das LLM bestimmt origin_tier aus dem Nachrichtentext "
    "(z.B. 'Kobaltmine' → Upstream, 'Zellfabrik-Brand' → Midstream-Cell)."
)
add_code("TierWeight  =  DISTANCE_WEIGHT[ |TIER_ORDER[facility] − TIER_ORDER[origin_tier]| ]")
add_heading("Stufen-Reihenfolge (TIER_ORDER):", 3)
add_code("Upstream = 0  →  Midstream-BGM = 1  →  Midstream-Cell = 2  →  Downstream = 3")
add_heading("Distanz-Gewichte:", 3)
add_table(
    headers=["Distanz", "TierWeight", "Bedeutung"],
    rows=[
        ["0", "1.0",  "Gleiche Stufe wie Ereignis → direkt betroffen"],
        ["1", "0.6",  "Eine Stufe entfernt"],
        ["2", "0.35", "Zwei Stufen entfernt"],
        ["3", "0.15", "Drei Stufen entfernt"],
    ],
    col_widths=[3, 3, 11],
)
add_heading("Beispiel S1 — Kobaltstreik (origin_tier = Upstream):", 3)
add_code(
    "Upstream       → Distanz 0 → TierWeight 1.0\n"
    "Midstream-BGM  → Distanz 1 → TierWeight 0.6\n"
    "Midstream-Cell → Distanz 2 → TierWeight 0.35\n"
    "Downstream     → Distanz 3 → TierWeight 0.15"
)
add_heading("Beispiel — Werksunfall Midstream-Cell (origin_tier = Midstream-Cell):", 3)
add_code(
    "Upstream       → Distanz 2 → TierWeight 0.35\n"
    "Midstream-BGM  → Distanz 1 → TierWeight 0.6\n"
    "Midstream-Cell → Distanz 0 → TierWeight 1.0   ← Unfallort\n"
    "Downstream     → Distanz 1 → TierWeight 0.6"
)
add_para("Quelle: origin_tier vom Risk Assessment Agent (LLM); segment-Feld aus NAATBatt.")

# ── Schritt 3: Vulnerability ─────────────────────────────────────────────────
add_heading("Schritt 3 — Vulnerability", 1)
add_code(
    "Vulnerability  =  0.30 × ImportDep\n"
    "               +  0.30 × SingleSource\n"
    "               +  0.25 × CapacityShare\n"
    "               +  0.15 × LeadTime"
)
add_para("Ergebnis liegt immer zwischen 0 und 1.")

add_heading("Die vier Felder im Detail", 2)

add_para(bold_parts=[("ImportDep", True), ("  (0 oder 1) — Ist das Material importabhängig?", False)])
add_para("  1 = Das Material kommt hauptsächlich aus dem Ausland → kein lokaler Ersatz möglich")
add_para("  0 = Es gibt lokale Alternativen")
add_para("  [simuliert — nach Materialtyp gesetzt]")

add_para(bold_parts=[("SingleSource", True), ("  (0 oder 1) — Gibt es nur wenige Anbieter weltweit?", False)])
add_para("  1 = ≤ 2 kommerzielle Anbieter in Nordamerika → kein Marktausgleich möglich")
add_para("  [simuliert — vorberechnet nach Materialtyp]")

add_para(bold_parts=[("CapacityShare", True), ("  (0.0 – 1.0) — Marktanteil dieser Fabrik", False)])
add_code("CapacityShare = Kapazität dieser Fabrik / Gesamtkapazität aller Fabriken (gleiches Material)")
add_para("  Beispiel: 12.000 MT / 85.000 MT = 0.14")

add_para(bold_parts=[("LeadTime", True), ("  (0.0 – 1.0) — Wie lange dauert Wiederbeschaffung?", False)])
add_code("LeadTime_norm = lead_time_weeks / 12   (gecappt auf 1.0)")
add_para("  [simuliert]: Upstream=12W, Midstream-BGM=8W, Midstream-Cell=6W, Downstream=4W")

add_heading("Warum diese Gewichte?", 2)
add_table(
    headers=["Feld", "Gewicht", "Begründung"],
    rows=[
        ["ImportDep",    "30 %", "Kein lokaler Ersatz = sofortiger Engpass"],
        ["SingleSource", "30 %", "Wenige Anbieter = kein Marktausgleich möglich"],
        ["CapacityShare","25 %", "Großer Marktanteil = großer systemischer Schaden"],
        ["LeadTime",     "15 %", "Lange Wartezeit verschlimmert kurzfristigen Engpass"],
    ],
    col_widths=[4, 3, 10],
)

# ── Schritt 4: ResilienceDiscount ─────────────────────────────────────────────
add_heading("Schritt 4 — ResilienceDiscount", 1)
add_para("Wenn es viele Ersatzlieferanten gibt, wird der RiskScore reduziert. "
         "Der Discount ist auf 0.5 begrenzt — auch mit Alternativen bleiben Umstellungszeit, "
         "Verträge und Qualitätsprüfung als Restrisiko.")
add_code(
    "AltCapacityRatio   =  Σ Kapazität Ersatzfabriken  /  Σ Kapazität betroffener Fabriken\n\n"
    "ResilienceDiscount =  min(AltCapacityRatio / 2,  0.5)"
)
add_table(
    headers=["AltCapacityRatio", "Discount", "Bedeutung"],
    rows=[
        ["0",      "0 %",        "Keine Alternativen — kein Abzug"],
        ["0.5",    "25 %",       "Ersatz deckt die Hälfte"],
        ["≥ 1.0",  "50 % (Max)", "Ausreichend Ersatz vorhanden"],
    ],
    col_widths=[4.5, 3.5, 9],
)

# ── Beispielrechnung ──────────────────────────────────────────────────────────
add_heading("Beispielrechnung — S1: Kobaltstreik DRC", 1)
add_table(
    headers=["Parameter", "Wert", "Herkunft"],
    rows=[
        ["Severity",      "4",    "LLM: monatelanger Streik, keine Ersatzlieferung"],
        ["TierWeight",    "1.0",  "Upstream-Mine"],
        ["ImportDep",     "1",    "Kobalt kommt aus dem Ausland [simuliert]"],
        ["SingleSource",  "1",    "≤ 2 Anbieter in Nordamerika [simuliert]"],
        ["CapacityShare", "0.14", "12.000 / 85.000 MT"],
        ["LeadTime",      "1.0",  "30 Wochen → gecappt auf 1.0 [simuliert]"],
    ],
    col_widths=[4, 2.5, 10.5],
)
add_code(
    "Vulnerability  = 0.30×1 + 0.30×1 + 0.25×0.14 + 0.15×1.0\n"
    "               = 0.30 + 0.30 + 0.035 + 0.15  =  0.785\n\n"
    "AltCapacityRatio    = 38.000 / 24.000  =  1.58\n"
    "ResilienceDiscount  = min(1.58 / 2, 0.5)  =  0.5\n\n"
    "RiskScore_roh       = 4 × 1.0 × 0.785 × (1 − 0.5)  =  1.57\n"
    "RiskScore (0–100)   = 1.57 / 5.0 × 100  =  31.4"
)

# ── Globale Kennzahlen ────────────────────────────────────────────────────────
add_heading("Globale Kennzahlen (separat vom RiskScore)", 1)
add_para("Diese Zahlen stehen im Berichts-Header. Sie zeigen das Gesamtbild des Ereignisses "
         "und werden nicht für das Ranking der einzelnen Fabriken verwendet.")
add_code(
    "BetroffeneKapazität%   =  Σ Kapazität betroffener Fabriken / Σ Gesamtkapazität × 100\n"
    "AlternativeKapazität%  =  Σ Kapazität der Ersatzfabriken   / Σ Gesamtkapazität × 100"
)
add_para("Beispiel S1:  BetroffeneKapazität% = 28 %   |   AlternativeKapazität% = 45 %")

# ── Klassifikation ────────────────────────────────────────────────────────────
add_heading("Risiko-Klassifikation", 1)
add_table(
    headers=["RiskScore", "Kategorie", "Empfehlung"],
    rows=[
        ["0 – 25",   "Niedrig",  "Beobachten"],
        ["25 – 50",  "Mittel",   "Proaktiv bewerten"],
        ["50 – 75",  "Hoch",     "Maßnahmen entwickeln"],
        ["75 – 100", "Kritisch", "Sofort handeln"],
    ],
    col_widths=[3, 3.5, 10.5],
    row_colors=["DCEDC8", "FFF9C4", "FFE0B2", "FFCDD2"],
)

# ── Überblick ─────────────────────────────────────────────────────────────────
add_heading("Überblick: Wer berechnet was?", 1)
add_table(
    headers=["Agent", "Output", "Methode"],
    rows=[
        ["Intake Agent",          "material, region, filtered_text",                           "LLM"],
        ["Risk Assessment Agent", "Severity (1–5), risk_type, reason",                          "LLM"],
        ["Network Agent",         "TierWeight, affected_nodes, alt_nodes",                      "Deterministisch (NetworkX)"],
        ["Data Retrieval Agent",  "Vulnerability-Felder, AltCapacityRatio, Kapazität%",         "Deterministisch (CSV + Arithmetik)"],
        ["Synthesis Agent",       "RiskScore (0–100), Top-3-Ranking, Risikobericht",            "LLM + Deterministisch"],
        ["Validation Agent",      "valid, failure_type (minor/severe), issues",                 "LLM"],
    ],
    col_widths=[4.5, 8, 4.5],
)

# ── Speichern ─────────────────────────────────────────────────────────────────
out = "/Users/yut/Projekt/sc_risk_system/docs/risk_model.docx"
doc.save(out)
print(f"Gespeichert: {out}")
