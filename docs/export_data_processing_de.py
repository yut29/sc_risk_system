"""
Erstellt eine einfache deutsche Word-Zusammenfassung der Datenverarbeitung.
Ausgabe: docs/Datenverarbeitung_Zusammenfassung.docx
"""

from pathlib import Path
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUTPUT = Path(__file__).parent / "Datenverarbeitung_Zusammenfassung.docx"

doc = Document()

# ── Titel ─────────────────────────────────────────────────────────────────────
title = doc.add_heading("Datenverarbeitung — SC Risk System", level=1)
title.alignment = WD_ALIGN_PARAGRAPH.LEFT

doc.add_paragraph(
    "Datenquelle: NAATBatt / NREL North American Lithium-Ion Battery Supply Chain "
    "Database, Stand März 2026. Verarbeitung in zwei Schritten: "
    "Anlagenbereinigung und Wissensgraphaufbau."
)
doc.add_paragraph()

# ── Datenklassifikation ───────────────────────────────────────────────────────
doc.add_heading("Datenklassifikation", level=2)
doc.add_paragraph(
    "Alle verwendeten Daten lassen sich in drei Kategorien einteilen:"
)

tbl_cls = doc.add_table(rows=4, cols=3)
tbl_cls.style = "Table Grid"
hdr = tbl_cls.rows[0].cells
hdr[0].text = "Typ"
hdr[1].text = "Deutsch"
hdr[2].text = "Betroffene Felder"
for i, (typ, de, felder) in enumerate([
    ("① Reale Daten",
     "Direkt aus NAATBatt, unverändert",
     "Unternehmensname, Adresse, Segment, Produkttyp, Koordinaten, "
     "production_capacity_raw (wo vorhanden), production_units"),
    ("② Simulierte Annahmen",
     "Fehlende Daten durch nachvollziehbare Methode approximiert",
     "Alle Kanten im Wissensgraphen (Keyword-Matching + geografische Nähe; "
     "NAATBatt enthält keine echten Lieferbeziehungen)"),
    ("③ Modellvereinfachungen",
     "Bewusste Designentscheidung mit expliziter Annahme",
     "lead_time_weeks (Stufenwerte), import_dependency (literaturbasierte Regel, USGS NIR), "
     "supplier_concentration (literaturbasierte Regel, USGS/IEA), import_origin_region (Lookup)"),
], 1):
    tbl_cls.rows[i].cells[0].text = typ
    tbl_cls.rows[i].cells[1].text = de
    tbl_cls.rows[i].cells[2].text = felder
doc.add_paragraph()

# ── Schritt 1 ─────────────────────────────────────────────────────────────────
doc.add_heading("Schritt 1 — Anlagenbereinigung", level=2)
doc.add_paragraph(
    "Ausgangspunkt: 1.280 Einträge in der Rohdatenbank "
    "(Excel, Tabellenblatt Append2)."
)

doc.add_heading("Filterregeln", level=3)
for s in [
    'Nur Anlagen mit Status "Commercial" (aktiv in Betrieb)',
    "Nur Nordamerika: USA, Kanada, Mexiko",
    "Nur relevante Lieferkettenstufen: Upstream, Midstream, Downstream",
    "Entfernung von 2 Anlagen ohne Koordinaten",
]:
    doc.add_paragraph(s, style="List Bullet")
doc.add_paragraph("Ergebnis nach Filterung: 386 Anlagen.")

# Midstream-Unterteilung
doc.add_heading("Midstream-Unterteilung", level=3)
doc.add_paragraph("Die Lieferkette für Batterien verläuft über vier Stufen:")
chain = doc.add_paragraph()
chain.add_run("Upstream → BGM (Battery Grade Materials) → Cell → Downstream").bold = True

doc.add_paragraph(
    "Die Rohdaten enthalten nur den Oberbegriff Midstream. "
    "Fuer die Risikomodellierung wird zwischen zwei Stufen unterschieden, "
    "da Risiken unterschiedlich schnell durch die Kette propagieren:"
)

tbl = doc.add_table(rows=3, cols=3)
tbl.style = "Table Grid"
hdr = tbl.rows[0].cells
hdr[0].text = "Stufe"
hdr[1].text = "Vollbezeichnung"
hdr[2].text = "Kriterium (Produktfeld)"
r1 = tbl.rows[1].cells
r1[0].text = "Midstream-BGM"
r1[1].text = "Battery Grade Materials"
r1[2].text = "Kein Zellfertigungs-Schlüsselwort"
r2 = tbl.rows[2].cells
r2[0].text = "Midstream-Cell"
r2[1].text = "Cell Manufacturing"
r2[2].text = "Enthält: cell, pouch, cylindrical, prismatic, cell assembly …"
doc.add_paragraph()

doc.add_heading("Abgeleitete Felder", level=3)
fields = [
    ("material_keywords",
     "Schlüsselwörter aus dem Produktfeld (cobalt, lithium, graphite …)"),
    ("capacity_source",
     '"naatbatt" = Originalwert aus NAATBatt (37 %); "unknown" = kein Wert vorhanden (63 %). '
     "Fehlende Kapazität wird nicht als Nullkapazität behandelt."),
    ("supplier_concentration",
     "True, wenn globaler Markt hochkonzentriert: cobalt (DRC 73 % Minenproduktion) "
     "und nmc/nca (China ≥ 95 % PCAM-Weltmarkt). Quellen: USGS MCS 2026, IEA GCMO 2025."),
    ("import_dependency",
     "True, wenn Segment ≠ Upstream UND Material nordamerikanisch nettoimportabhängig "
     "(USGS Net Import Reliance 2026: Co 79 %, Graphit 100 %, Mn 100 %, Li > 50 %). "
     "Upstream-Anlagen immer False (sie sind selbst Produktionsstandorte)."),
    ("import_origin_region",
     "Herkunftsregion nach Material (z. B. Kobalt → Afrika / DRC)"),
    ("lead_time_weeks",
     "Modellvereinfachung nach Stufe: Upstream 12 Wo., BGM 8 Wo., Cell 6 Wo., Downstream 4 Wo."),
]
for name, desc in fields:
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(name + ": ").bold = True
    p.add_run(desc)
doc.add_paragraph()

# ── Schritt 2 ─────────────────────────────────────────────────────────────────
doc.add_heading("Schritt 2 — Wissensgraph", level=2)
doc.add_paragraph(
    "Die NAATBatt-Datenbank enthält keine echten Lieferbeziehungen. "
    "Alle 1.537 Kanten wurden auf Basis von Materialschlüsselwörtern und "
    "geografischer Nähe simuliert (② Simulierte Annahmen)."
)

doc.add_heading("Kantengenerierung", level=3)
edge_rows = [
    ("Upstream → BGM",
     "Materialschlüsselwort-Matching\n"
     "Beispiel: cobalt → NMC-Kathode"),
    ("BGM → Cell",
     "Schlüsselwort-Matching + geografische Nähe\n"
     "Jede BGM-Anlage wird mit den K = 6 nächstgelegenen passenden Cell-Anlagen verbunden"),
    ("Cell → Downstream",
     "Nur geografische Nähe\n"
     "Jede Cell-Anlage wird mit den K = 5 nächstgelegenen Downstream-Anlagen verbunden "
     "+ 3 zufällige Fernverbindungen (reproduzierbar per facility_id als Seed)"),
]
tbl_e = doc.add_table(rows=len(edge_rows)+1, cols=2)
tbl_e.style = "Table Grid"
tbl_e.rows[0].cells[0].text = "Verbindungsschicht"
tbl_e.rows[0].cells[1].text = "Logik"
for i, (layer, logic) in enumerate(edge_rows, 1):
    tbl_e.rows[i].cells[0].text = layer
    tbl_e.rows[i].cells[1].text = logic
doc.add_paragraph()

# ── Ergebnisse ────────────────────────────────────────────────────────────────
doc.add_heading("Ergebnisse", level=2)

doc.add_heading("Anlagen nach Lieferkettenstufe", level=3)
tbl2 = doc.add_table(rows=6, cols=3)
tbl2.style = "Table Grid"
tbl2.rows[0].cells[0].text = "Stufe"
tbl2.rows[0].cells[1].text = "Sep 2025 (alt)"
tbl2.rows[0].cells[2].text = "März 2026 (neu)"
for i, (s, old, new) in enumerate([
    ("Upstream", "26", "29"),
    ("Midstream-BGM (Battery Grade Materials)", "42", "114"),
    ("Midstream-Cell (Cell Manufacturing)", "44", "54"),
    ("Downstream", "173", "189"),
    ("Gesamt", "285", "386"),
], 1):
    tbl2.rows[i].cells[0].text = s
    tbl2.rows[i].cells[1].text = old
    tbl2.rows[i].cells[2].text = new
doc.add_paragraph()

doc.add_heading("Wissensgraph", level=3)
tbl_g = doc.add_table(rows=6, cols=2)
tbl_g.style = "Table Grid"
tbl_g.rows[0].cells[0].text = "Kennzahl"
tbl_g.rows[0].cells[1].text = "Wert"
for i, (k, v) in enumerate([
    ("Knoten gesamt", "386"),
    ("Kanten gesamt", "1.537 (alle simuliert)"),
    ("Upstream → BGM", "421 Kanten"),
    ("BGM → Cell", "684 Kanten"),
    ("Cell → Downstream", "432 Kanten (4 % Konnektivität; zuvor 81 % mit altem Länder-Ansatz)"),
], 1):
    tbl_g.rows[i].cells[0].text = k
    tbl_g.rows[i].cells[1].text = v
doc.add_paragraph()

doc.add_heading("Kapazitätsdaten", level=3)
tbl3 = doc.add_table(rows=3, cols=3)
tbl3.style = "Table Grid"
tbl3.rows[0].cells[0].text = "Status"
tbl3.rows[0].cells[1].text = "Anzahl"
tbl3.rows[0].cells[2].text = "Anteil"
for i, (s, n, p) in enumerate([
    ("NAATBatt-Originalwert (capacity_known = True)", "142", "37 %"),
    ("Kein Wert in NAATBatt (capacity_known = False)", "244", "63 %"),
], 1):
    tbl3.rows[i].cells[0].text = s
    tbl3.rows[i].cells[1].text = n
    tbl3.rows[i].cells[2].text = p
doc.add_paragraph(
    "Hinweis: Fehlende Kapazität ≠ Nullkapazität. "
    "CapacityShare und AltCapacityRatio werden nur berechnet, wenn capacity_known = True."
)
doc.add_paragraph()

doc.add_heading("Materialabdeckung", level=3)
tbl4 = doc.add_table(rows=6, cols=2)
tbl4.style = "Table Grid"
tbl4.rows[0].cells[0].text = "Material"
tbl4.rows[0].cells[1].text = "Anlagen"
for i, (m, n) in enumerate([
    ("Cobalt", "39"), ("Lithium", "39"), ("Nickel", "36"),
    ("Graphite", "23"), ("Manganese", "19"),
], 1):
    tbl4.rows[i].cells[0].text = m
    tbl4.rows[i].cells[1].text = n
doc.add_paragraph()

# ── Einschränkungen ───────────────────────────────────────────────────────────
doc.add_heading("Einschränkungen", level=2)
for l in [
    "Keine echten Lieferbeziehungen: Alle Kanten sind simuliert, nicht verifiziert.",
    "63 % der Anlagen ohne Kapazitätsdaten: BetroffeneKapazität % ist eine konservative Untergrenze.",
    "Nur Nordamerika abgedeckt: Indirekte Abhängigkeiten von Asien (v. a. China) nicht erfasst.",
    "Midstream-BGM-Kapazitäten nicht vergleichbar (MT / mm² / GWh / L gemischt) — CapacityShare nur für Upstream.",
]:
    doc.add_paragraph(l, style="List Bullet")

doc.add_paragraph()
doc.add_paragraph(
    "Stand: 2026-07-06  |  Bearbeitung: Yutong Liu  |  "
    "Betreuer: Baris Albayrak (FAPS, FAU)"
)

doc.save(OUTPUT)
print(f"Gespeichert: {OUTPUT}")
