# -*- coding: utf-8 -*-
"""
Erstellt eine deutsche Word-Zusammenfassung des Risikomodells.
Ausgabe: docs/Risikomodell_Zusammenfassung.docx
"""

from pathlib import Path
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUTPUT = Path(__file__).parent / "Risikomodell_Zusammenfassung.docx"

doc = Document()

# ── Titel ─────────────────────────────────────────────────────────────────────
title = doc.add_heading("Risikomodell — SC Risk System", level=1)
title.alignment = WD_ALIGN_PARAGRAPH.LEFT

doc.add_paragraph(
    "Das Modell berechnet für jede betroffene Anlage einen RiskScore (0-100). "
    "Ein hoher Score bedeutet: Diese Anlage ist stark gefährdet und es gibt kaum Alternativen. "
    "Die drei Anlagen mit dem höchsten Score erscheinen im Risikobericht."
)
doc.add_paragraph()

# ── Hauptformel ───────────────────────────────────────────────────────────────
doc.add_heading("Hauptformel", level=2)
p = doc.add_paragraph()
p.add_run("RiskScore = Severity * TierWeight * Vulnerability * (1 − ResilienceDiscount)").bold = True
doc.add_paragraph("Normierung: RiskScore (0-100) = RiskScore_roh / 5,0 * 100")
doc.add_paragraph(
    "Das theoretische Maximum beträgt 5,0 "
    "(Severity = 5, TierWeight = 1,0, Vulnerability = 1,0, kein Discount)."
)
doc.add_paragraph()

# ── Severity ──────────────────────────────────────────────────────────────────
doc.add_heading("Schritt 1 — Severity (Schweregrad des Ereignisses)", level=2)
doc.add_paragraph("Quelle: LLM liest den Nachrichtentext und gibt einen Wert von 1-5 zurück.")
tbl_sev = doc.add_table(rows=6, cols=2)
tbl_sev.style = "Table Grid"
tbl_sev.rows[0].cells[0].text = "Wert"
tbl_sev.rows[0].cells[1].text = "Bedeutung"
for i, (v, b) in enumerate([
    ("1", "Sehr gering — kleiner lokaler Streik, nur 1 Tag"),
    ("2", "Gering — kurze Lieferverzögerung"),
    ("3", "Mittel — monatelanger Streik"),
    ("4", "Hoch — Exportverbot für kritisches Material"),
    ("5", "Kritisch — kompletter Lieferstopp, keine Alternative"),
], 1):
    tbl_sev.rows[i].cells[0].text = v
    tbl_sev.rows[i].cells[1].text = b
doc.add_paragraph()

# ── TierWeight ────────────────────────────────────────────────────────────────
doc.add_heading("Schritt 2 — TierWeight (Nähe zum Ereignisursprung)", level=2)
doc.add_paragraph(
    "TierWeight hängt davon ab, wie weit eine Anlage vom Ereignisursprung (origin_tier) entfernt ist. "
    "Das LLM bestimmt origin_tier aus dem Nachrichtentext (z. B. Kobaltmine -> Upstream)."
)
p2 = doc.add_paragraph()
p2.add_run("TierWeight = DISTANCE_WEIGHT[ |TIER_ORDER[facility] − TIER_ORDER[origin_tier]| ]").bold = True

tbl_tw = doc.add_table(rows=5, cols=2)
tbl_tw.style = "Table Grid"
tbl_tw.rows[0].cells[0].text = "Distanz zum Ereignisursprung"
tbl_tw.rows[0].cells[1].text = "TierWeight"
for i, (d, w) in enumerate([
    ("0 (gleiche Stufe)", "1,0"),
    ("1", "0,6"),
    ("2", "0,35"),
    ("3", "0,15"),
], 1):
    tbl_tw.rows[i].cells[0].text = d
    tbl_tw.rows[i].cells[1].text = w
doc.add_paragraph()

# ── Vulnerability ─────────────────────────────────────────────────────────────
doc.add_heading("Schritt 3 — Vulnerability (Verwundbarkeit der Anlage)", level=2)
p3 = doc.add_paragraph()
p3.add_run(
    "Vulnerability = 0,30 * ImportDep + 0,30 * SupplierConcentration "
    "+ 0,25 * CapacityShare + 0,15 * LeadTime"
).bold = True
doc.add_paragraph()

# Übersichtstabelle
tbl_vuln = doc.add_table(rows=5, cols=3)
tbl_vuln.style = "Table Grid"
tbl_vuln.rows[0].cells[0].text = "Feld"
tbl_vuln.rows[0].cells[1].text = "Gewicht"
tbl_vuln.rows[0].cells[2].text = "Kurzbeschreibung"
for i, (f, w, d) in enumerate([
    ("ImportDep",            "30 %", "Geopolitisches Risiko — Nordamerika nettoimportabhaengig?"),
    ("SupplierConcentration","30 %", "Marktstrukturrisiko — wenige globale Anbieter?"),
    ("CapacityShare",        "25 %", "Systemischer Marktanteil der Anlage (nur Upstream)"),
    ("LeadTime",             "15 %", "Normierte Wiederbeschaffungszeit nach Stufe"),
], 1):
    tbl_vuln.rows[i].cells[0].text = f
    tbl_vuln.rows[i].cells[1].text = w
    tbl_vuln.rows[i].cells[2].text = d
doc.add_paragraph()

# ── ImportDep ──────────────────────────────────────────────────────────────────
doc.add_heading("ImportDep — Geopolitisches Risiko (Gewicht: 30 %)", level=3)
doc.add_paragraph(
    "Gibt an, ob Nordamerika bei diesem Material nettoimportabhaengig ist "
    "UND die Anlage kein Produktionsstandort ist."
)
doc.add_paragraph(
    "Segment bezeichnet die Lieferkettenstufe einer Anlage: "
    "Upstream (Rohstoffproduktion), Midstream-BGM (Batteriematerialien), "
    "Midstream-Cell (Zellproduktion), Downstream (Endprodukte)."
)
doc.add_paragraph(
    "Regel: 1 wenn Segment != Upstream UND Material in USGS-Importliste. "
    "Upstream-Anlagen erhalten immer 0 — sie sind selbst Produzenten, keine Importeure."
)
tbl_imp = doc.add_table(rows=6, cols=3)
tbl_imp.style = "Table Grid"
tbl_imp.rows[0].cells[0].text = "Material"
tbl_imp.rows[0].cells[1].text = "Net Import Reliance"
tbl_imp.rows[0].cells[2].text = "ImportDep"
for i, (m, nir, v) in enumerate([
    ("Cobalt",  "79 %",      "1"),
    ("Graphit", "100 %",     "1"),
    ("Mangan",  "100 %",     "1"),
    ("Lithium", "> 50 %",    "1"),
    ("Nickel",  "ca. 100 %", "1"),
], 1):
    tbl_imp.rows[i].cells[0].text = m
    tbl_imp.rows[i].cells[1].text = nir
    tbl_imp.rows[i].cells[2].text = v
doc.add_paragraph("Quelle: USGS Mineral Commodity Summaries 2026.")
doc.add_paragraph()

# ── SupplierConcentration ──────────────────────────────────────────────────────
doc.add_heading("SupplierConcentration — Marktstrukturrisiko (Gewicht: 30 %)", level=3)
doc.add_paragraph(
    "Gibt an, ob der globale Markt fuer dieses Material von wenigen Anbietern dominiert wird, "
    "sodass ein Ausfall kaum kompensiert werden kann."
)
tbl_sc = doc.add_table(rows=4, cols=3)
tbl_sc.style = "Table Grid"
tbl_sc.rows[0].cells[0].text = "Material"
tbl_sc.rows[0].cells[1].text = "Begruendung"
tbl_sc.rows[0].cells[2].text = "Wert"
for i, (m, b, v) in enumerate([
    ("cobalt",
     "DRC 73 % + Indonesia 14 % der Weltminenproduktion; 2-3 Konzerne dominieren. Quelle: USGS MCS 2026.",
     "1"),
    ("nmc / nca",
     "China >= 95 % PCAM-Weltmarktanteil; < 5 Nicht-China-Grosslieferanten. Quelle: IEA GCMO 2025.",
     "1"),
    ("lithium, graphite, nickel, manganese",
     "10+ Minenunternehmen weltweit; kein einzelner Anbieter dominiert.",
     "0"),
], 1):
    tbl_sc.rows[i].cells[0].text = m
    tbl_sc.rows[i].cells[1].text = b
    tbl_sc.rows[i].cells[2].text = v
doc.add_paragraph()

# ── CapacityShare ──────────────────────────────────────────────────────────────
doc.add_heading("CapacityShare — Systemischer Marktanteil (Gewicht: 25 %)", level=3)
doc.add_paragraph(
    "Anteil dieser Anlage an der gesamten NAATBatt-Upstream-Kapazitaet "
    "fuer dasselbe Material (nur Upstream, MT/yr einheitlich und vergleichbar)."
)
p_cs = doc.add_paragraph()
p_cs.add_run(
    "CapacityShare = Kapazitaet dieser Anlage / Σ Upstream-Kapazitaet (gleiches Material, Nordamerika)"
).bold = True
doc.add_paragraph(
    "Bedingungen:\n"
    "- Nur berechnet wenn capacity_known = True (NAATBatt-Originalwert vorhanden).\n"
    "- Nur fuer Upstream-Anlagen (MT/yr vergleichbar); Midstream und Downstream = 0.\n"
    "- Bei capacity_known = False: CapacityShare = 0 (konservativ; fehlende Daten != keine Kapazitaet)."
)
doc.add_paragraph()

# ── LeadTime ──────────────────────────────────────────────────────────────────
doc.add_heading("LeadTime — Normierte Wiederbeschaffungszeit (Gewicht: 15 %)", level=3)
doc.add_paragraph(
    "Modellvereinfachung: feste Vorlaufzeit je Lieferkettenstufe, "
    "normiert auf eine maximale Referenzzeit von 12 Wochen (Upstream)."
)
p_lt = doc.add_paragraph()
p_lt.add_run("LeadTime = lead_time_weeks / 12   (auf 1,0 begrenzt)").bold = True
tbl_lt = doc.add_table(rows=5, cols=3)
tbl_lt.style = "Table Grid"
tbl_lt.rows[0].cells[0].text = "Segment"
tbl_lt.rows[0].cells[1].text = "Wochen"
tbl_lt.rows[0].cells[2].text = "LeadTime"
for i, (s, w, v) in enumerate([
    ("Upstream",       "12", "1,00"),
    ("Midstream-BGM",  "8",  "0,67"),
    ("Midstream-Cell", "6",  "0,50"),
    ("Downstream",     "4",  "0,33"),
], 1):
    tbl_lt.rows[i].cells[0].text = s
    tbl_lt.rows[i].cells[1].text = w
    tbl_lt.rows[i].cells[2].text = v
doc.add_paragraph()

# ── ResilienceDiscount ────────────────────────────────────────────────────────
doc.add_heading("Schritt 4 — ResilienceDiscount (Verfügbarkeit von Alternativen)", level=2)
doc.add_paragraph(
    "AltCapacityRatio = Σ Kapazität der Ersatzanlagen / Kapazität der betroffenen Anlage"
)
p4 = doc.add_paragraph()
p4.add_run("ResilienceDiscount = min(AltCapacityRatio / 2,  0,5)").bold = True
doc.add_paragraph(
    "Der Discount ist auf 0,5 begrenzt — auch bei guten Alternativen bleibt "
    "ein Restrisiko (Umstellungszeit, Verträge, Qualitätsprüfung). "
    "Nur berechnet wenn capacity_known = True; sonst Discount = 0 (konservativ)."
)
doc.add_paragraph()

# ── Kapazitätsdatenstatus ─────────────────────────────────────────────────────
doc.add_heading("Kapazitätsdatenqualität", level=2)
tbl_cap = doc.add_table(rows=3, cols=3)
tbl_cap.style = "Table Grid"
tbl_cap.rows[0].cells[0].text = "capacity_source"
tbl_cap.rows[0].cells[1].text = "capacity_known"
tbl_cap.rows[0].cells[2].text = "Behandlung im Modell"
for i, (src, known, treat) in enumerate([
    ("naatbatt", "True", "Echter NAATBatt-Wert — fließt in CapacityShare und AltCapacityRatio ein"),
    ("unknown", "False",
     "Datenlücke, nicht Nullkapazität — CapacityShare = 0, ResilienceDiscount = 0 (konservativ)"),
], 1):
    tbl_cap.rows[i].cells[0].text = src
    tbl_cap.rows[i].cells[1].text = known
    tbl_cap.rows[i].cells[2].text = treat
doc.add_paragraph()

# ── Beispielrechnung ──────────────────────────────────────────────────────────
doc.add_heading("Beispielrechnung — S1: Kobaltstreik DRC", level=2)
tbl_ex = doc.add_table(rows=7, cols=3)
tbl_ex.style = "Table Grid"
tbl_ex.rows[0].cells[0].text = "Parameter"
tbl_ex.rows[0].cells[1].text = "Wert"
tbl_ex.rows[0].cells[2].text = "Herkunft"
for i, (p, v, h) in enumerate([
    ("Severity", "4", "LLM: monatelanger Streik"),
    ("TierWeight", "1,0", "Upstream-Mine, Distanz 0"),
    ("ImportDep", "1", "Co NIR = 79 % (USGS MCS 2026)"),
    ("SupplierConcentration", "1", "DRC 73 % + Indonesia 14 % (USGS MCS 2026)"),
    ("CapacityShare", "0,14", "12.000 / 85.000 MT"),
    ("LeadTime", "1,0", "30 Wochen -> gecappt"),
], 1):
    tbl_ex.rows[i].cells[0].text = p
    tbl_ex.rows[i].cells[1].text = v
    tbl_ex.rows[i].cells[2].text = h

doc.add_paragraph()
doc.add_paragraph(
    "Vulnerability = 0,30*1 + 0,30*1 + 0,25*0,14 + 0,15*1,0 = 0,785\n"
    "AltCapacityRatio = 38.000 / 24.000 = 1,58  ->  ResilienceDiscount = min(0,79; 0,5) = 0,5\n"
    "RiskScore_roh = 4 * 1,0 * 0,785 * 0,5 = 1,57\n"
    "RiskScore (0-100) = 1,57 / 5,0 * 100 = 31,4"
)
doc.add_paragraph()

# ── Risikokategorien ──────────────────────────────────────────────────────────
doc.add_heading("Risiko-Klassifikation", level=2)
tbl_risk = doc.add_table(rows=5, cols=3)
tbl_risk.style = "Table Grid"
tbl_risk.rows[0].cells[0].text = "RiskScore"
tbl_risk.rows[0].cells[1].text = "Kategorie"
tbl_risk.rows[0].cells[2].text = "Empfehlung"
for i, (r, k, e) in enumerate([
    ("0 - 25", "Niedrig", "Beobachten"),
    ("25 - 50", "Mittel", "Proaktiv bewerten"),
    ("50 - 75", "Hoch", "Maßnahmen entwickeln"),
    ("75 - 100", "Kritisch", "Sofort handeln"),
], 1):
    tbl_risk.rows[i].cells[0].text = r
    tbl_risk.rows[i].cells[1].text = k
    tbl_risk.rows[i].cells[2].text = e

doc.add_paragraph()
doc.add_paragraph(
    "Stand: 2026-07-06  |  Bearbeitung: Yutong Liu  |  "
    "Betreuer: Baris Albayrak (FAPS, FAU)"
)

doc.save(OUTPUT)
print(f"Gespeichert: {OUTPUT}")
